"""Output generation module for transcription pipeline."""

import re
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict


class OutputGenerator:
    """Generates markdown and Word document outputs from formatted transcripts."""

    def __init__(self, output_dir: Path):
        """Initialize with output directory.

        Args:
            output_dir: Base directory for all outputs
        """
        self.output_dir = Path(output_dir)
        self.transcripts_dir = self.output_dir / "transcripts"
        self.docs_dir = self.output_dir / "docs"

        # Create output directories
        self.transcripts_dir.mkdir(parents=True, exist_ok=True)
        self.docs_dir.mkdir(parents=True, exist_ok=True)
    
    def get_user_docs_dir(self, user_subdir: Optional[str] = None) -> Path:
        """Get the docs directory for a specific user.
        
        Args:
            user_subdir: Optional subdirectory for user-specific outputs (e.g., "keira")
            
        Returns:
            Path to the docs directory (user-specific or default)
        """
        if user_subdir:
            user_dir = self.docs_dir / user_subdir
            user_dir.mkdir(parents=True, exist_ok=True)
            return user_dir
        return self.docs_dir

    def generate_outputs(
        self, formatted_text: str, note_type: str, filename: str, metadata: dict = None
    ) -> dict:
        """Generate appropriate output formats based on note_type.

        Args:
            formatted_text: The formatted transcript content
            note_type: Type of note (meeting, supervision, client, lecture, braindump)
            filename: Original filename (used for deriving title)
            metadata: Optional metadata dictionary with duration, speakers, etc.

        Returns:
            Dictionary with markdown_path, docx_path, and title
        """
        metadata = metadata or {}
        title = self._derive_title(filename, note_type)

        # Determine which formats to generate
        generate_md = note_type in ("meeting", "supervision", "client", "braindump")
        generate_docx = note_type in ("meeting", "supervision", "client", "lecture")

        markdown_path: Optional[Path] = None
        docx_path: Optional[Path] = None

        # Generate markdown if needed
        if generate_md:
            md_content = self._create_markdown(formatted_text, title, metadata)
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            markdown_path = self.transcripts_dir / f"{safe_title}.md"
            markdown_path.write_text(md_content, encoding='utf-8')

        # Generate docx if needed
        if generate_docx:
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            docx_path = self.docs_dir / f"{safe_title}.docx"

            # Try pandoc first, fallback to python-docx
            if self._pandoc_available() and markdown_path:
                self._create_docx_with_pandoc(markdown_path, docx_path)
            else:
                self._create_docx(formatted_text, title, docx_path, metadata)

        return {
            "markdown_path": markdown_path,
            "docx_path": docx_path,
            "title": title,
        }

    def generate_multi_stage_output(
        self,
        content: str,
        filename_base: str,
        suffix: str,
        stage_name: str,
        metadata: dict = None,
        generate_docx: bool = True,
        docs_dir: Optional[Path] = None
    ) -> List[Dict]:
        """Generate output file for a multi-stage processing stage.

        Args:
            content: The content to save
            filename_base: Base filename (without extension)
            suffix: Suffix to add to filename (e.g., "_filtered", "_clean")
            stage_name: Name of the stage (for logging/metadata)
            metadata: Optional metadata dictionary
            generate_docx: Whether to also generate Word document (default: True)
            docs_dir: Optional custom docs directory (defaults to self.docs_dir)

        Returns:
            List of output file info dicts with 'path', 'type', 'stage' keys
        """
        metadata = metadata or {}
        outputs = []
        
        # Use custom docs directory if provided, otherwise use default
        target_docs_dir = docs_dir if docs_dir else self.docs_dir

        # Generate filename with suffix
        md_filename = self._derive_filename(filename_base, suffix, ".md")
        md_path = self.transcripts_dir / md_filename

        # Create markdown with YAML frontmatter for stage files
        md_content = self._create_stage_markdown(content, stage_name, metadata)
        
        # Write markdown file
        md_path.write_text(md_content, encoding='utf-8')
        
        outputs.append({
            "path": md_path,
            "type": "markdown",
            "stage": stage_name,
        })

        # Generate Word document if requested
        if generate_docx:
            docx_filename = self._derive_filename(filename_base, suffix, ".docx")
            docx_path = target_docs_dir / docx_filename
            
            # Create title from filename and stage
            title = self._derive_title_from_filename(filename_base, suffix)
            
            try:
                # Try pandoc first (better formatting), fallback to python-docx
                if self._pandoc_available():
                    self._create_docx_with_pandoc(md_path, docx_path)
                else:
                    self._create_docx(content, title, docx_path, metadata)
                
                outputs.append({
                    "path": docx_path,
                    "type": "docx",
                    "stage": stage_name,
                })
            except Exception as e:
                # Log error but don't fail - markdown is the primary output
                import logging
                logging.getLogger(__name__).warning(
                    f"Failed to generate .docx for stage {stage_name}: {e}"
                )

        return outputs

    def _derive_filename(self, base: str, suffix: str, extension: str) -> str:
        """Derive a filename with suffix.
        
        Args:
            base: Base filename (without extension)
            suffix: Suffix to add (e.g., "_filtered")
            extension: File extension (including dot)
        
        Returns:
            Complete filename
        """
        # Clean base name
        safe_base = re.sub(r'[^\w\s-]', '', base).strip().replace(' ', '_')
        
        # Add suffix if provided
        if suffix:
            # Remove leading underscore if present in suffix
            suffix = suffix.lstrip('_')
            return f"{safe_base}_{suffix}{extension}"
        else:
            return f"{safe_base}{extension}"

    def _derive_title_from_filename(self, filename_base: str, suffix: str) -> str:
        """Derive a clean title from filename base and stage suffix.
        
        Args:
            filename_base: Original filename without extension
            suffix: Stage suffix (e.g., "_filtered", "_clean")
        
        Returns:
            Clean title for document
        """
        # Remove timestamp prefix (common formats)
        name = re.sub(r'^(\d{8}_\d{6}[_-]?|\d{4}-\d{2}-\d{2}-\d{2}-\d{2}-\d{2}-?)', '', filename_base)
        
        # Replace underscores and hyphens with spaces
        name = name.replace('_', ' ').replace('-', ' ')
        
        # Clean up multiple spaces
        name = re.sub(r'\s+', ' ', name).strip()
        
        # Capitalize each word
        words = name.split()
        capitalized = [word.capitalize() for word in words]
        title = ' '.join(capitalized)
        
        # Add stage suffix as subtitle
        if suffix:
            stage_name = suffix.lstrip('_').replace('_', ' ').title()
            title = f"{title} ({stage_name})"
        
        return title

    def _create_markdown(self, content: str, title: str, metadata: dict) -> str:
        """Create markdown content with YAML frontmatter.

        Args:
            content: The formatted transcript content
            title: The derived title
            metadata: Metadata dictionary

        Returns:
            Markdown string with YAML frontmatter
        """
        lines = ["---"]
        # Quote title to handle special YAML characters like colons
        lines.append(f'title: "{title.replace(chr(34), chr(92)+chr(34))}"')
        lines.append(f"date: {datetime.now().strftime('%Y-%m-%d')}")

        if "type" in metadata:
            lines.append(f"type: {metadata['type']}")

        if "duration" in metadata:
            lines.append(f"duration: {metadata['duration']}")

        if "speakers" in metadata:
            if isinstance(metadata["speakers"], list):
                speakers_str = ", ".join(metadata["speakers"])
            else:
                speakers_str = str(metadata["speakers"])
            lines.append(f"speakers: {speakers_str}")

        # Add any other metadata fields
        for key, value in metadata.items():
            if key not in ("type", "duration", "speakers"):
                lines.append(f"{key}: {value}")

        lines.append("---")
        lines.append("")
        lines.append(content)

        return "\n".join(lines)

    def _create_stage_markdown(self, content: str, stage_name: str, metadata: dict) -> str:
        """Create markdown content for a multi-stage output with YAML frontmatter.

        Args:
            content: The stage output content
            stage_name: Name of the processing stage
            metadata: Metadata dictionary

        Returns:
            Markdown string with YAML frontmatter
        """
        lines = ["---"]
        lines.append(f"stage: {stage_name}")
        lines.append(f"date: {datetime.now().strftime('%Y-%m-%d')}")
        lines.append(f"processed_at: {datetime.now().isoformat()}")
        
        if "profile" in metadata:
            lines.append(f"profile: {metadata['profile']}")
        
        if "duration" in metadata:
            lines.append(f"audio_duration: {metadata['duration']}")

        # Add any other metadata fields
        for key, value in metadata.items():
            if key not in ("stage", "duration", "profile"):
                lines.append(f"{key}: {value}")

        lines.append("---")
        lines.append("")
        lines.append(content)

        return "\n".join(lines)

    def _create_docx(
        self, content: str, title: str, output_path: Path, metadata: dict = None
    ) -> None:
        """Create Word document using python-docx.

        Args:
            content: The formatted transcript content
            title: The document title
            output_path: Path to save the document
            metadata: Optional metadata dictionary
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )

        doc = Document()

        # Add title page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata to title page
        if metadata:
            doc.add_paragraph()  # Spacing

            if "date" in metadata:
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(f"Date: {metadata['date']}")
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if "duration" in metadata:
                duration_para = doc.add_paragraph()
                duration_run = duration_para.add_run(f"Duration: {metadata['duration']}")
                duration_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if "speakers" in metadata:
                speakers = metadata["speakers"]
                if isinstance(speakers, list):
                    speakers_str = ", ".join(speakers)
                else:
                    speakers_str = str(speakers)
                speakers_para = doc.add_paragraph()
                speakers_run = speakers_para.add_run(f"Speakers: {speakers_str}")
                speakers_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page break after title page
        doc.add_page_break()

        # Parse and convert content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                # Empty line - add spacing
                i += 1
                continue

            # Check for headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                # python-docx has Heading 1-9, map # -> 1, ## -> 2, etc.
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue

            # Check for bullet lists
            bullet_match = re.match(r'^[\*\-\+]\s+(.+)$', line)
            if bullet_match:
                text = bullet_match.group(1)
                doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # Check for numbered lists
            numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if numbered_match:
                text = numbered_match.group(2)
                doc.add_paragraph(text, style='List Number')
                i += 1
                continue

            # Regular paragraph
            # Handle inline formatting (bold, italic)
            para = doc.add_paragraph()
            self._add_formatted_text(para, line)

            i += 1

        # Save document
        doc.save(str(output_path))

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with inline markdown formatting to a paragraph.

        Args:
            paragraph: docx paragraph object
            text: Text possibly containing markdown formatting
        """
        # Pattern for bold and italic
        # Process **bold** and *italic* and `code`
        pattern = r'(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)'

        parts = re.split(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code/monospace
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Regular text
                paragraph.add_run(part)

    def _create_docx_with_pandoc(self, markdown_path: Path, output_path: Path) -> None:
        """Create Word document using pandoc.

        Args:
            markdown_path: Path to markdown file
            output_path: Path to save the docx file
        """
        try:
            subprocess.run(
                ["pandoc", str(markdown_path), "-o", str(output_path), "-f", "markdown", "-t", "docx"],
                check=True,
                capture_output=True,
                text=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr}") from e

    def _pandoc_available(self) -> bool:
        """Check if pandoc is available in the system."""
        try:
            subprocess.run(
                ["pandoc", "--version"],
                check=True,
                capture_output=True,
            )
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def _derive_title(self, filename: str, note_type: str) -> str:
        """Derive a clean title from filename.

        Args:
            filename: Original filename
            note_type: Type of note

        Returns:
            Clean title with note type prefix if needed
        """
        # Remove extension
        name = Path(filename).stem

        # Remove timestamp prefix (common formats: 20240115_143022_, 2024-01-15-14-30-22-)
        name = re.sub(r'^(\d{8}_\d{6}[_-]?|\d{4}-\d{2}-\d{2}-\d{2}-\d{2}-\d{2}-?)', '', name)

        # Replace underscores and hyphens with spaces
        name = name.replace('_', ' ').replace('-', ' ')

        # Clean up multiple spaces
        name = re.sub(r'\s+', ' ', name).strip()

        # Capitalize each word
        words = name.split()
        capitalized = [word.capitalize() for word in words]
        title = ' '.join(capitalized)

        # Add note type prefix if not already present (case-insensitive check)
        note_type_cap = note_type.capitalize()
        if not re.search(rf'\b{re.escape(note_type_cap)}\b', title, re.IGNORECASE):
            title = f"{note_type_cap}: {title}"

        return title
